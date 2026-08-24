import io
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.models.resume import ResumeDocument


class DOCXService:
    def generate_docx(self, resume: ResumeDocument) -> bytes:
        """
        Renders a structured ResumeDocument into a genuine Microsoft Word (.docx) binary document.
        """
        doc = docx.Document()

        # Set 0.5 inch margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        p = resume.personal_info
        f = resume.formatting

        # Header - Name
        name_p = doc.add_paragraph()
        name_run = name_p.add_run(f"{p.firstName} {p.lastName}".upper())
        name_run.font.name = "Arial"
        name_run.font.size = Pt(20)
        name_run.bold = True
        name_run.font.color.rgb = RGBColor(17, 24, 39)
        name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_p.paragraph_format.space_after = Pt(2)

        # Contact Info Line
        contact_parts = [item for item in [p.email, p.phone, p.location] if item]
        if p.linkedin:
            contact_parts.append(p.linkedin)
        if p.github:
            contact_parts.append(p.github)
        if p.portfolio:
            contact_parts.append(p.portfolio)

        contact_p = doc.add_paragraph()
        contact_run = contact_p.add_run(" • ".join(contact_parts))
        contact_run.font.name = "Arial"
        contact_run.font.size = Pt(9.5)
        contact_run.font.color.rgb = RGBColor(75, 85, 99)
        contact_p.paragraph_format.space_after = Pt(12)

        # Helper to add section headers
        def add_section_header(title_text):
            h_p = doc.add_paragraph()
            h_p.paragraph_format.space_before = Pt(10)
            h_p.paragraph_format.space_after = Pt(4)
            run = h_p.add_run(title_text.upper())
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(26, 115, 232)  # Knora Blue Accent

        # Summary
        if resume.summary and resume.summary.strip():
            add_section_header("Professional Summary")
            sum_p = doc.add_paragraph()
            sum_run = sum_p.add_run(resume.summary.strip())
            sum_run.font.name = "Arial"
            sum_run.font.size = Pt(10)
            sum_p.paragraph_format.space_after = Pt(8)

        # Education
        if resume.education:
            add_section_header("Education")
            for edu in resume.education:
                edu_p = doc.add_paragraph()
                edu_p.paragraph_format.space_after = Pt(2)
                
                degree_run = edu_p.add_run(f"{edu.degree} {f'in {edu.field}' if edu.field else ''}")
                degree_run.bold = True
                degree_run.font.size = Pt(10)
                
                dates = f" ({edu.startDate or ''} - {'Present' if edu.currentlyStudying else (edu.endDate or '')})"
                date_run = edu_p.add_run(dates)
                date_run.font.size = Pt(9.5)
                date_run.italic = True
                
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.space_after = Pt(6)
                inst_run = sub_p.add_run(f"{edu.institution}{f', {edu.location}' if edu.location else ''}")
                inst_run.font.size = Pt(9.5)
                if edu.grade:
                    grade_run = sub_p.add_run(f" | Grade: {edu.grade}")
                    grade_run.font.size = Pt(9.5)
                    grade_run.bold = True

        # Experience
        if resume.experience:
            add_section_header("Experience")
            for exp in resume.experience:
                exp_p = doc.add_paragraph()
                exp_p.paragraph_format.space_after = Pt(1)
                
                pos_run = exp_p.add_run(exp.position)
                pos_run.bold = True
                pos_run.font.size = Pt(10)
                
                dates = f" ({exp.startDate or ''} - {'Present' if exp.currentlyWorking else (exp.endDate or '')})"
                date_run = exp_p.add_run(dates)
                date_run.font.size = Pt(9.5)
                date_run.italic = True
                
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.space_after = Pt(3)
                comp_run = sub_p.add_run(f"{exp.company}{f', {exp.location}' if exp.location else ''}")
                comp_run.font.size = Pt(9.5)
                comp_run.bold = True

                if exp.description:
                    desc_p = doc.add_paragraph()
                    desc_p.paragraph_format.space_after = Pt(3)
                    d_run = desc_p.add_run(exp.description)
                    d_run.font.size = Pt(9.5)

                if exp.achievements:
                    for ach in exp.achievements:
                        if ach.strip():
                            b_p = doc.add_paragraph(style='List Bullet')
                            b_p.paragraph_format.space_after = Pt(2)
                            b_run = b_p.add_run(ach.strip())
                            b_run.font.size = Pt(9.5)

        # Projects
        if resume.projects:
            add_section_header("Projects")
            for proj in resume.projects:
                proj_p = doc.add_paragraph()
                proj_p.paragraph_format.space_after = Pt(1)
                
                t_run = proj_p.add_run(proj.title)
                t_run.bold = True
                t_run.font.size = Pt(10)
                
                if proj.startDate:
                    dates = f" ({proj.startDate} - {proj.endDate or ''})"
                    d_run = proj_p.add_run(dates)
                    d_run.font.size = Pt(9)
                    d_run.italic = True
                    
                if proj.technologies:
                    sub_p = doc.add_paragraph()
                    sub_p.paragraph_format.space_after = Pt(2)
                    tech_run = sub_p.add_run(f"Technologies: {proj.technologies}")
                    tech_run.font.size = Pt(9)
                    tech_run.italic = True
                    
                if proj.description:
                    desc_p = doc.add_paragraph()
                    desc_p.paragraph_format.space_after = Pt(6)
                    desc_run = desc_p.add_run(proj.description)
                    desc_run.font.size = Pt(9.5)

        # Skills
        if resume.skills:
            add_section_header("Skills & Technologies")
            categories = {}
            for sk in resume.skills:
                cat = sk.category or "Technical"
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(sk.name)

            for cat, sk_list in categories.items():
                sk_p = doc.add_paragraph()
                sk_p.paragraph_format.space_after = Pt(3)
                cat_run = sk_p.add_run(f"{cat}: ")
                cat_run.bold = True
                cat_run.font.size = Pt(9.5)
                list_run = sk_p.add_run(", ".join(sk_list))
                list_run.font.size = Pt(9.5)

        # Certifications
        if resume.certifications:
            add_section_header("Certifications")
            for cert in resume.certifications:
                c_p = doc.add_paragraph()
                c_p.paragraph_format.space_after = Pt(3)
                name_run = c_p.add_run(cert.name)
                name_run.bold = True
                name_run.font.size = Pt(9.5)
                if cert.issuer:
                    iss_run = c_p.add_run(f" — {cert.issuer}")
                    iss_run.font.size = Pt(9.5)

        # Languages
        if resume.languages:
            add_section_header("Languages")
            lang_p = doc.add_paragraph()
            lang_p.paragraph_format.space_after = Pt(6)
            lang_strs = [f"{l.language} ({l.proficiency})" for l in resume.languages if l.language]
            l_run = lang_p.add_run(" • ".join(lang_strs))
            l_run.font.size = Pt(9.5)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()


docx_service = DOCXService()
